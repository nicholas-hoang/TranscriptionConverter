import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import app


def docx_to_txt(file):
    """
    Converts a Word document to a DataFrame.

    Args:
        file (str): The path to the Word document.

    Returns:
        pandas.DataFrame: The Word document converted to a DataFrame.

    """
    document = Document(file)
    doc = []
    for para in document.paragraphs:
        doc.append(para.text)
    df = pd.DataFrame(doc)
    return df


def format_table(dataframe):
    """
    Formats the input DataFrame into a table with columns for the timestamp, speaker, and text.

    Args:
        df (pandas.DataFrame): The input DataFrame containing the transcript.

    Returns:
        pandas.DataFrame: The input DataFrame formatted into a table with columns for the timestamp, speaker, and text.

    """
    df = dataframe
    # Split the transcripts into a list of strings
    df = df[0].str.split("\n")
    df = df.apply(pd.Series)
    # Rename the columns
    df.rename(columns={0: "TimeStamp", 1: "Speaker", 2: "Text"}, inplace=True)
    # Split the timestamp into start and end times
    df.TimeStamp = df["TimeStamp"].apply(lambda x: x.split(" --> "))
    df["Start"] = df["TimeStamp"].apply(lambda x: x[0])
    df["End"] = df["TimeStamp"].apply(lambda x: x[1])

    return df


def map_speakers(dataframe):
    """
    Maps each speaker to a unique integer label.

    Args:
        df (pandas.DataFrame): The input DataFrame containing the speaker column.

    Returns:
        pandas.DataFrame: The input DataFrame with a new column containing the integer labels for each speaker.
        dict: A dictionary mapping each speaker to a unique integer label.

    """
    df = dataframe
    speakers_dict = {
        speaker: index for index, speaker in enumerate(df.Speaker.unique())
    }
    df["Label"] = df.Speaker.map(speakers_dict)
    return df, speakers_dict


# def redact_speakers_input(dataframe, speakers_dictionary):

#     # Key Values Pairs will be kept as Name: Index for the time being,
#     # until I can figure out how to make it {Index: Name}

#     print('Here are the speakers in this transcript: ')
#     for key, value in speakers_dictionary.items():
#         print(f'{key}: {value}')

#     # Ask the user which speaker they want to redact
#     speaker_to_redact = input('Which speaker would you like to redact? ')


def concatenate_text_with_timestamp_and_speaker_by_label(dataframe):
    """
    Concatenates the text, start timestamp, end timestamp, and speaker name for each group of rows with the same label
    in the input DataFrame.

    Args:
        df (pandas.DataFrame): The input DataFrame containing the text, timestamp, and speaker columns.

    Returns:
        list: A list of dictionaries, where each dictionary contains the concatenated text, start timestamp, end timestamp,
        and speaker name for a group of rows with the same label.
    """
    df = dataframe
    concatenated_data = []
    current_text = ""
    start_timestamp = None
    speaker = None

    for index, row in df.iterrows():
        if index == 0 or row["Label"] == df.at[index - 1, "Label"]:
            if not start_timestamp:
                start_timestamp = row["TimeStamp"][0]
            if not speaker:
                speaker = row["Speaker"]
            current_text += " " + row["Text"]
            end_timestamp = row["TimeStamp"][1]
        else:
            concatenated_data.append(
                {
                    "text": current_text.strip(),
                    "start_timestamp": start_timestamp,
                    "end_timestamp": end_timestamp,
                    "speaker": speaker,
                }
            )
            current_text = row["Text"]
            start_timestamp = row["TimeStamp"][0]
            end_timestamp = row["TimeStamp"][1]
            speaker = row["Speaker"]

    concatenated_data.append(
        {
            "text": current_text.strip(),
            "start_timestamp": start_timestamp,
            "end_timestamp": end_timestamp,
            "speaker": speaker,
        }
    )
    return concatenated_data


def write_to_word_doc(concatenated_data, filename):
    """
    Writes the concatenated text, start timestamp, end timestamp, and speaker name for each group of rows to a Word document.

    Args:
        concatenated_data (list): A list of dictionaries, where each dictionary contains the concatenated text, start timestamp,
        end timestamp, and speaker name for a group of rows with the same label.
        filename (str): The name of the file to save the Word document as.

    Returns:
        None
    """
    # Create a new Word document
    document = Document()

    # Add a heading to the document
    document.add_heading("Meeting Transcription", level=0)

    # Loop through the concatenated data and add each group to the document
    for group in concatenated_data:
        # Add the start and end timestamps to the document
        document.add_paragraph(
            f"[{group['start_timestamp']} - {group['end_timestamp']}]"
        )

        # Add the speaker name to the document and make it bold.
        speaker_paragraph = document.add_paragraph()
        speaker_paragraph.add_run(f"Speaker: {group['speaker']}").bold = True

        # Add the concatenated text to the document
        document.add_paragraph(group["text"])

    # Save the document
    document.save(filename)


def app(file_path):
    df = docx_to_txt(file_path)
    df = format_table(df)
    df, speakers_dictionary = map_speakers(df)
    # redact_speakers_input(df, speakers_dictionary)
    concatenate_data = concatenate_text_with_timestamp_and_speaker_by_label(df)
    write_to_word_doc(concatenate_data, "data/concatenated_text.docx")


if __name__ == "__main__":
    pass
