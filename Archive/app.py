import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import os
from docx import Document
import pandas as pd
import os
from os.path import expanduser
import datetime


class FileUploader(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Audio Transcription Formatter for Rapid Analysis by. Nick Hoang")
        self.geometry("600x400")
        self.setup_gui()

    def setup_gui(self):
        # # Set the style of the GUI
        # style = ttk.Style(self)
        # self.tk.call("source", "forest-light.tcl")
        # # self.tk.call("source", "forest-dark.tcl")
        # style.theme_use("forest-light")

        # Left Frame for File Selection
        left_frame = tk.Frame(self)
        left_frame.pack(side="left", fill="both", padx=20, pady=20)

        # File Selection Label and Entry
        select_label = tk.Label(
            left_frame, text="Select a File:", font=("Helvetica", 16)
        )
        select_label.grid(row=0, column=0, sticky="w", pady=(0, 10))

        self.file_path_entry = tk.Entry(left_frame, font=("Helvetica", 12), width=60)
        self.file_path_entry.grid(row=0, column=1, padx=(10, 0), pady=(0, 10))

        # Buttons (centered on the same row)
        select_button = tk.Button(
            left_frame, text="Browse", command=self.select_file, font=("Helvetica", 14)
        )
        select_button.grid(row=1, column=0, pady=(0, 10), sticky="e")

        self.run_button = tk.Button(
            left_frame,
            text="Run",
            font=("Helvetica", 14),
            state="disabled",
            command=self.run_module,
        )
        self.run_button.grid(row=1, column=1, padx=(10, 0), pady=(0, 10), sticky="w")

        # Feedback Label
        self.feedback_label = tk.Label(
            left_frame, text="", font=("Helvetica", 12), fg="green"
        )
        self.feedback_label.grid(row=3, column=0, columnspan=2)

    def select_file(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            self.file_path_entry.delete(0, tk.END)  # Clear previous text
            self.file_path_entry.insert(0, file_path)  # Insert selected file path
            self.run_button.config(state="active")  # Enable the run button

    def run_module(self):
        file_path = self.file_path_entry.get()  # Get the selected file path
        if file_path and os.path.exists(file_path):
            converter = TranscriptionConverter(file_path)
            converter.convert_and_write()
            self.feedback_label.config(
                text="File processing completed successfully!\n\nFormatted Transcript saved in Downloads folder ヾ( ˃ᴗ˂ )◞ • *✰.",
                fg="green",
            )
        else:
            self.feedback_label.config(
                text="Error: Invalid file path or file does not exist. (¬_¬')", fg="red"
            )


class TranscriptionConverter:
    def __init__(self, file_path):
        self.file_path = file_path

    def docx_to_txt(self):
        document = Document(self.file_path)
        doc = [para.text for para in document.paragraphs]
        df = pd.DataFrame(doc, columns=["Text"])
        return df

    def format_table(self, dataframe):
        # Split the transcripts into a list of strings
        df = dataframe["Text"].str.split("\n", expand=True)
        # Rename the columns
        df.columns = ["TimeStamp", "Speaker", "Text"]
        # Split the timestamp into start and end times
        df[["Start", "End"]] = df["TimeStamp"].str.split(" --> ", expand=True)
        return df

    def map_speakers(self, dataframe):
        speakers_dict = {
            speaker: index
            for index, speaker in enumerate(dataframe["Speaker"].unique())
        }
        dataframe["Label"] = dataframe["Speaker"].map(speakers_dict)
        return dataframe, speakers_dict

    def concatenate_text_with_timestamp_and_speaker_by_label(self, dataframe):
        concatenated_data = []
        current_text = ""
        start_timestamp = None
        speaker = None

        for index, row in dataframe.iterrows():
            if (
                start_timestamp is None
                or row["Label"] == dataframe.at[index - 1, "Label"]
            ):
                if start_timestamp is None:
                    start_timestamp = row["Start"]
                if speaker is None:
                    speaker = row["Speaker"]
                current_text += " " + row["Text"]
                end_timestamp = row["End"]
            else:
                concatenated_data.append(
                    {
                        "text": current_text.strip(),
                        "start_timestamp": start_timestamp,
                        "end_timestamp": end_timestamp,
                        "speaker": speaker,
                    }
                )
                current_text = row["Text"]
                start_timestamp = row["Start"]
                end_timestamp = row["End"]
                speaker = row["Speaker"]

        concatenated_data.append(
            {
                "text": current_text.strip(),
                "start_timestamp": start_timestamp,
                "end_timestamp": end_timestamp,
                "speaker": speaker,
            }
        )
        return concatenated_data

    def write_to_word_doc(self, concatenated_data, filename):
        document = Document()
        document.add_heading("Meeting Transcription", level=0)

        for group in concatenated_data:
            document.add_paragraph(
                f"[{group['start_timestamp']} - {group['end_timestamp']}]"
            )
            speaker_paragraph = document.add_paragraph()
            speaker_paragraph.add_run(f"Speaker: {group['speaker']}").bold = True
            document.add_paragraph(group["text"])

        document.save(filename)

    def convert_and_write(self):
        # Get the user's Downloads folder path
        downloads_folder = expanduser("~") + "/Downloads"
        # Create the complete file path in the Downloads folder

        # Extract the original file name without the extension
        original_file_name = os.path.splitext(os.path.basename(self.file_path))[0]

        # Append the desired suffix and date to the original file name
        output_file_name = f"{original_file_name}-CLEANED-{datetime.datetime.now().strftime('%Y-%m-%d')}"

        # Create the complete file path in the Downloads folder
        output_file_path = os.path.join(downloads_folder, output_file_name + ".docx")

        df = self.docx_to_txt()
        df = self.format_table(df)
        df, speakers_dictionary = self.map_speakers(df)
        concatenated_data = self.concatenate_text_with_timestamp_and_speaker_by_label(
            df
        )
        self.write_to_word_doc(concatenated_data, output_file_path)


if __name__ == "__main__":
    app = FileUploader()
    app.mainloop()
