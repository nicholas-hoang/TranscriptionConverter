import os
from os.path import expanduser
import gradio as gr
from docx import Document
import pandas as pd
import datetime

"""
gradioapp.py
------------
This script provides a Gradio interface for processing transcriptions in DOCX format.
It includes a class `TranscriptionConverter` that handles the conversion of the transcription
from DOCX to a cleaned format, and a main function that sets up and launches the Gradio interface.

Classes:
--------
TranscriptionConverter: Handles the conversion of transcriptions from DOCX to a cleaned format.

Functions:
----------
process_transcription(file): Processes the uploaded transcription file and returns the path of the cleaned output file.

main(): Sets up and launches the Gradio interface.

Usage:
------
This script is intended to be run as a standalone script. It will launch a Gradio interface
that allows the user to upload a DOCX file, which will then be processed and cleaned.
The cleaned transcription can then be downloaded through the Gradio interface.
"""


class TranscriptionConverter:
    def __init__(self, file_path):
        self.file_path = file_path

    def docx_to_txt(self):
        document = Document(self.file_path)
        doc = [para.text for para in document.paragraphs]
        df = pd.DataFrame(doc, columns=["Text"])
        return df

    def format_table(self, dataframe):
        # Split the transcripts into a list of strings
        df = dataframe["Text"].str.split("\n", expand=True)
        # Rename the columns
        df.columns = ["TimeStamp", "Speaker", "Text"]
        # Split the timestamp into start and end times
        df[["Start", "End"]] = df["TimeStamp"].str.split(" --> ", expand=True)
        return df

    def map_speakers(self, dataframe):
        speakers_dict = {
            speaker: index
            for index, speaker in enumerate(dataframe["Speaker"].unique())
        }
        dataframe["Label"] = dataframe["Speaker"].map(speakers_dict)
        return dataframe, speakers_dict

    def concatenate_text_with_timestamp_and_speaker_by_label(self, dataframe):
        concatenated_data = []
        current_text = ""
        start_timestamp = None
        speaker = None

        for index, row in dataframe.iterrows():
            if (
                start_timestamp is None
                or row["Label"] == dataframe.at[index - 1, "Label"]
            ):
                if start_timestamp is None:
                    start_timestamp = row["Start"]
                if speaker is None:
                    speaker = row["Speaker"]
                current_text += " " + row["Text"]
                end_timestamp = row["End"]
            else:
                concatenated_data.append(
                    {
                        "text": current_text.strip(),
                        "start_timestamp": start_timestamp,
                        "end_timestamp": end_timestamp,
                        "speaker": speaker,
                    }
                )
                current_text = row["Text"]
                start_timestamp = row["Start"]
                end_timestamp = row["End"]
                speaker = row["Speaker"]

        concatenated_data.append(
            {
                "text": current_text.strip(),
                "start_timestamp": start_timestamp,
                "end_timestamp": end_timestamp,
                "speaker": speaker,
            }
        )
        return concatenated_data

    def write_to_word_doc(self, concatenated_data, filename):
        document = Document()
        document.add_heading("Meeting Transcription", level=0)

        for group in concatenated_data:
            document.add_paragraph(
                f"[{group['start_timestamp']} - {group['end_timestamp']}]"
            )
            speaker_paragraph = document.add_paragraph()
            speaker_paragraph.add_run(f"Speaker: {group['speaker']}").bold = True
            document.add_paragraph(group["text"])

        document.save(filename)

    def convert_and_write(self):
        # Extract the original file name without the extension
        original_file_name = os.path.splitext(os.path.basename(self.file_path))[0]

        # Append the desired suffix and date to the original file name
        output_file_name = f"{original_file_name}-CLEANED-{datetime.datetime.now().strftime('%Y-%m-%d')}"

        df = self.docx_to_txt()
        df = self.format_table(df)
        df, speakers_dictionary = self.map_speakers(df)
        concatenated_data = self.concatenate_text_with_timestamp_and_speaker_by_label(
            df
        )
        self.write_to_word_doc(concatenated_data, output_file_name)


def process_transcription(file):
    # Gradio automatically saves the uploaded file, so we can directly use its path
    file_path = file.name
    converter = TranscriptionConverter(file_path)
    df = converter.docx_to_txt()
    df = converter.format_table(df)
    df, speakers_dictionary = converter.map_speakers(df)
    concatenated_data = converter.concatenate_text_with_timestamp_and_speaker_by_label(
        df
    )
    output_file_name = (
        f"{file_path}_CLEANED_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"
    )

    converter.write_to_word_doc(concatenated_data, output_file_name)
    return output_file_name


def main():
    iface = gr.Interface(
        fn=process_transcription,
        inputs=gr.File(label="Upload your DOCX file"),
        outputs=gr.File(label="Download cleaned transcript"),
    )
    iface.launch()


if __name__ == "__main__":
    main()
